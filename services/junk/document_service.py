"""
Document Service - File processing and web scraping
Handles PDF, Word, text files and website content extraction
"""

import os
import logging
import asyncio
import aiofiles
from typing import Dict, List, Optional, Any, BinaryIO
from datetime import datetime
import hashlib

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import requests
from playwright.async_api import async_playwright

from utils.logger_config import get_logger

logger = get_logger(__name__)

class DocumentService:
    def __init__(self):
        self.supported_types = {
            "application/pdf": {"extension": ".pdf", "processor": "_process_pdf"},
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {
                "extension": ".docx", "processor": "_process_docx"
            },
            "application/msword": {"extension": ".doc", "processor": "_process_docx"},
            "text/plain": {"extension": ".txt", "processor": "_process_text"},
            "text/html": {"extension": ".html", "processor": "_process_html"},
            "application/json": {"extension": ".json", "processor": "_process_json"}
        }
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.max_pages = 100  # Max pages for PDF
        self.website_timeout = 30000  # 30 seconds

    async def initialize(self):
        """Initialize document service"""
        logger.info("✅ Document Service initialized")

    async def process_document(self, file: Any, user_id: str, metadata: Dict = None) -> Dict[str, Any]:
        """Process uploaded document file"""
        try:
            # Read file content
            content = await file.read()
            filename = file.filename
            content_type = file.content_type

            logger.info(f"Processing document: {filename} ({content_type}) for user {user_id}")

            # Validate file type
            if content_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {content_type}")

            # Validate file size
            if len(content) > self.max_file_size:
                raise ValueError(f"File size exceeds limit of {self.max_file_size} bytes")

            # Process file based on type
            file_info = self.supported_types[content_type]
            processor_method = getattr(self, file_info["processor"])
            extracted_content = await processor_method(content, filename)

            # Prepare document data
            document_data = {
                "title": filename,
                "content": extracted_content["text"],
                "source": "file_upload",
                "document_type": file_info["extension"],
                "file_size": len(content),
                "metadata": {
                    **(metadata or {}),
                    "original_filename": filename,
                    "mime_type": content_type,
                    "processed_at": datetime.utcnow().isoformat(),
                    "page_count": extracted_content.get("page_count", 1),
                    "word_count": len(extracted_content["text"].split()),
                    "processing_time": extracted_content.get("processing_time", 0)
                }
            }

            # Process through vector store
            from services.vector_store_service import VectorStoreService
            vector_service = VectorStoreService()
            vector_results = await vector_service.process_document(document_data, user_id)

            return {
                "success": True,
                "document": document_data,
                "vector_chunks": len(vector_results),
                "processing_details": extracted_content
            }

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "document": {"title": getattr(file, "filename", "unknown")}
            }

    async def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF file"""
        start_time = datetime.now()

        try:
            import io
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from all pages (up to max_pages)
            text_parts = []
            page_count = min(len(pdf_reader.pages), self.max_pages)

            for i in range(page_count):
                page = pdf_reader.pages[i]
                text_parts.append(page.extract_text())

            full_text = "\n".join(text_parts)
            processing_time = (datetime.now() - start_time).total_seconds()

            return {
                "text": full_text,
                "page_count": page_count,
                "processing_time": processing_time,
                "metadata": {
                    "total_pages": len(pdf_reader.pages),
                    "processed_pages": page_count
                }
            }

        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise Exception(f"PDF processing failed: {str(e)}")

    async def _process_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word document"""
        start_time = datetime.now()

        try:
            import io
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            full_text = "\n".join(text_parts)
            processing_time = (datetime.now() - start_time).total_seconds()

            return {
                "text": full_text,
                "page_count": 1,  # Word doesn't have explicit pages in text
                "processing_time": processing_time,
                "metadata": {
                    "paragraph_count": len(doc.paragraphs)
                }
            }

        except Exception as e:
            logger.error(f"Word document processing failed: {e}")
            raise Exception(f"Word processing failed: {str(e)}")

    async def _process_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text file"""
        start_time = datetime.now()

        try:
            # Try different encodings
            text = None
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise Exception("Could not decode text file")

            processing_time = (datetime.now() - start_time).total_seconds()

            return {
                "text": text,
                "page_count": 1,
                "processing_time": processing_time,
                "metadata": {
                    "encoding": encoding,
                    "line_count": len(text.split("\n"))
                }
            }

        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise Exception(f"Text processing failed: {str(e)}")

    async def _process_html(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process HTML file"""
        start_time = datetime.now()

        try:
            html_content = content.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            # Extract text
            text = soup.get_text()
            # Clean up whitespace
            text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

            processing_time = (datetime.now() - start_time).total_seconds()

            # Extract metadata
            title = soup.find('title')
            title_text = title.get_text() if title else filename

            meta_description = soup.find('meta', attrs={'name': 'description'})
            description = meta_description.get('content') if meta_description else ""

            return {
                "text": text,
                "page_count": 1,
                "processing_time": processing_time,
                "metadata": {
                    "title": title_text,
                    "description": description,
                    "html_elements": len(soup.find_all())
                }
            }

        except Exception as e:
            logger.error(f"HTML processing failed: {e}")
            raise Exception(f"HTML processing failed: {str(e)}")

    async def _process_json(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process JSON file"""
        start_time = datetime.now()

        try:
            import json
            json_data = json.loads(content.decode('utf-8'))

            # Convert JSON to readable text
            text = self._json_to_text(json_data)
            processing_time = (datetime.now() - start_time).total_seconds()

            return {
                "text": text,
                "page_count": 1,
                "processing_time": processing_time,
                "metadata": {
                    "json_keys": list(json_data.keys()) if isinstance(json_data, dict) else [],
                    "data_type": type(json_data).__name__
                }
            }

        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            raise Exception(f"JSON processing failed: {str(e)}")

    def _json_to_text(self, obj: Any, depth: int = 0) -> str:
        """Convert JSON object to readable text"""
        if depth > 5:  # Prevent infinite recursion
            return "[deeply nested object]"

        if isinstance(obj, dict):
            parts = []
            for key, value in obj.items():
                value_text = self._json_to_text(value, depth + 1)
                parts.append(f"{key}: {value_text}")
            return "\n".join(parts)

        elif isinstance(obj, list):
            parts = []
            for i, item in enumerate(obj[:10]):  # Limit to first 10 items
                item_text = self._json_to_text(item, depth + 1)
                parts.append(f"Item {i+1}: {item_text}")
            if len(obj) > 10:
                parts.append(f"... and {len(obj) - 10} more items")
            return "\n".join(parts)

        else:
            return str(obj)

    async def scrape_website(self, url: str, user_id: str, method: str = "static") -> Dict[str, Any]:
        """Scrape website content"""
        try:
            logger.info(f"Scraping website: {url} using {method} method for user {user_id}")

            # Validate URL
            if not self._is_valid_url(url):
                raise ValueError("Invalid URL format")

            # Choose scraping method
            if method == "static":
                content = await self._scrape_static(url)
            elif method == "dynamic":
                content = await self._scrape_dynamic(url)
            else:
                raise ValueError(f"Unknown scraping method: {method}")

            # Prepare website data
            website_data = {
                "url": url,
                "title": content["title"],
                "content": content["text"],
                "metadata": {
                    "scraped_at": datetime.utcnow().isoformat(),
                    "method": method,
                    "content_length": len(content["text"]),
                    "word_count": len(content["text"].split()),
                    "status_code": content.get("status_code", 200)
                }
            }

            # Process through vector store
            from services.vector_store_service import VectorStoreService
            vector_service = VectorStoreService()
            vector_results = await vector_service.process_website(website_data, user_id)

            return {
                "success": True,
                "website": website_data,
                "vector_chunks": len(vector_results),
                "scraping_details": content
            }

        except Exception as e:
            logger.error(f"Website scraping failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "website": {"url": url}
            }

    async def _scrape_static(self, url: str) -> Dict[str, Any]:
        """Scrape website using requests (static content)"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }

            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # Remove unwanted elements
            for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                element.extract()

            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "Untitled"

            # Extract main content
            # Try to find main content area
            main_content = soup.find('main') or soup.find('article') or soup.find('div', {'id': 'content'}) or soup.find('body')

            text = main_content.get_text() if main_content else soup.get_text()
            # Clean up text
            text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

            return {
                "title": title_text,
                "text": text,
                "method": "static",
                "status_code": response.status_code
            }

        except Exception as e:
            raise Exception(f"Static scraping failed: {str(e)}")

    async def _scrape_dynamic(self, url: str) -> Dict[str, Any]:
        """Scrape website using Playwright (dynamic content)"""
        try:
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page()

                # Set user agent and viewport
                await page.set_user_agent('Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
                await page.set_viewport_size({"width": 1920, "height": 1080})

                # Navigate to page
                response = await page.goto(url, timeout=self.website_timeout)

                # Wait for content to load
                await page.wait_for_timeout(3000)

                # Extract content
                title = await page.title()

                # Remove unwanted elements
                await page.evaluate("""() => {
                    const elements = document.querySelectorAll('script, style, nav, footer, header, aside, .ads, .advertisement');
                    elements.forEach(el => el.remove());
                }""")

                # Get text content
                text = await page.evaluate('() => document.body.innerText')

                await browser.close()

                return {
                    "title": title or "Untitled",
                    "text": text.strip(),
                    "method": "dynamic",
                    "status_code": response.status if response else 200
                }

        except Exception as e:
            raise Exception(f"Dynamic scraping failed: {str(e)}")

    def _is_valid_url(self, url: str) -> bool:
        """Validate URL format"""
        from urllib.parse import urlparse
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc]) and result.scheme in ['http', 'https']
        except:
            return False

    async def health_check(self) -> Dict[str, Any]:
        """Health check for document service"""
        try:
            # Test basic functionality
            test_content = b"Test document content"
            result = await self._process_text(test_content, "test.txt")

            return {
                "status": "healthy",
                "supported_types": len(self.supported_types),
                "max_file_size": self.max_file_size,
                "test_processing": "passed" if result["text"] else "failed"
            }

        except Exception as e:
            logger.error(f"Document service health check failed: {e}")
            return {
                "status": "unhealthy",
                "error": str(e)
            }
